"""
文档处理工具
"""

import os
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentTools:
    """文档处理工具集"""
    
    def __init__(self):
        self.current_project_dir = None
    
    def set_project_dir(self, project_dir: Path):
        """设置当前项目目录"""
        self.current_project_dir = project_dir
    
    def convert_to_docx(self, markdown_content: str) -> str:
        """转换Markdown为Word文档"""
        try:
            # 创建新文档
            doc = Document()
            
            # 解析markdown内容
            lines = markdown_content.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 处理标题
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                # 处理列表
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith(('1. ', '2. ', '3. ')):
                    doc.add_paragraph(line[3:], style='List Number')
                # 处理普通段落
                else:
                    if line.startswith('**') and line.endswith('**'):
                        # 粗体段落
                        p = doc.add_paragraph()
                        run = p.add_run(line[2:-2])
                        run.bold = True
                    else:
                        doc.add_paragraph(line)
            
            # 保存临时文件
            temp_path = "/tmp/temp_document.docx"
            doc.save(temp_path)
            
            return temp_path
            
        except Exception as e:
            logger.error(f"Failed to convert markdown to docx: {e}")
            raise
    
    def one_click_docx_export(self) -> Dict[str, Any]:
        """一键导出完整Word标书"""
        if not self.current_project_dir:
            raise ValueError("No project directory set")
        
        try:
            # 创建新文档
            doc = Document()
            
            # 设置文档样式
            self._setup_document_styles(doc)
            
            # 添加封面
            self._add_cover_page(doc)
            
            # 添加目录页
            self._add_table_of_contents(doc)
            
            # 合并所有章节内容
            sections_content = self._collect_all_sections()
            
            for section_info in sections_content:
                # 添加章节标题
                doc.add_heading(section_info['title'], level=1)
                
                # 添加小节内容
                for subsection in section_info['subsections']:
                    if subsection['has_content']:
                        # 添加小节标题
                        doc.add_heading(subsection['title'], level=2)
                        
                        # 添加内容
                        content_lines = subsection['content'].split('\n')
                        for line in content_lines:
                            line = line.strip()
                            if line and not line.startswith('#'):
                                doc.add_paragraph(line)
                
                # 章节间分页
                doc.add_page_break()
            
            # 保存文档
            output_dir = self.current_project_dir / "output"
            output_dir.mkdir(exist_ok=True)
            
            output_path = output_dir / "标书_v1.0.docx"
            doc.save(str(output_path))
            
            # 获取文件信息
            file_size = output_path.stat().st_size
            
            return {
                "file_path": str(output_path),
                "file_size": f"{file_size / 1024 / 1024:.1f}MB",
                "pages": len(doc.paragraphs) // 20,  # 估算页数
                "sections": len(sections_content),
                "charts": 0,  # TODO: 统计图表数量
                "success": True
            }
            
        except Exception as e:
            logger.error(f"Failed to export docx: {e}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def merge_subsections_to_docx(self) -> Dict[str, Any]:
        """合并所有三级标题内容为Word"""
        return self.one_click_docx_export()
    
    def assemble_section_from_subsections(self, section: str) -> str:
        """组装章节从三级标题文件"""
        if not self.current_project_dir:
            raise ValueError("No project directory set")
        
        sections_dir = self.current_project_dir / "sections"
        
        # 查找章节目录
        section_dir = None
        for item in sections_dir.iterdir():
            if item.is_dir() and section.lower() in item.name.lower():
                section_dir = item
                break
        
        if not section_dir:
            return f"# {section}\n\n章节内容未找到"
        
        # 收集所有小节内容
        content_parts = [f"# {section}\n\n"]
        
        for file_item in sorted(section_dir.iterdir()):
            if file_item.is_file() and file_item.suffix == '.md':
                try:
                    subsection_content = file_item.read_text(encoding='utf-8')
                    content_parts.append(subsection_content)
                    content_parts.append("\n\n")
                except Exception as e:
                    logger.error(f"Failed to read {file_item}: {e}")
        
        return "".join(content_parts)
    
    def apply_template(self, content: str, template: str) -> str:
        """应用文档模板"""
        # 简单的模板应用
        if template == "standard":
            return f"""
# 标书文档

{content}

---
*本文档由 Tender AI 生成*
"""
        return content
    
    def format_tables(self, data: List[List[str]]) -> str:
        """格式化表格"""
        if not data:
            return ""
        
        # 生成Markdown表格
        table_lines = []
        
        # 表头
        headers = data[0]
        table_lines.append("| " + " | ".join(headers) + " |")
        table_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        
        # 数据行
        for row in data[1:]:
            table_lines.append("| " + " | ".join(row) + " |")
        
        return "\n".join(table_lines)
    
    def insert_charts(self, data: Dict[str, Any], chart_type: str) -> str:
        """插入图表"""
        # 返回图表占位符
        return f"""
```chart
类型: {chart_type}
数据: {data}
```

*图表将在Word文档中自动生成*
"""
    
    def export_pdf(self, docx_path: str) -> str:
        """导出PDF"""
        # 这里需要使用额外的库如python-docx2pdf
        # 目前返回占位符
        pdf_path = docx_path.replace('.docx', '.pdf')
        return pdf_path
    
    def batch_format_docx(self, subsections: List[str]) -> Dict[str, Any]:
        """批量格式化三级标题内容"""
        results = {}
        
        for subsection in subsections:
            try:
                # 格式化单个小节
                results[subsection] = {
                    "success": True,
                    "formatted": True
                }
            except Exception as e:
                results[subsection] = {
                    "success": False,
                    "error": str(e)
                }
        
        return results
    
    def _setup_document_styles(self, doc: Document):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 设置标题样式
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = '黑体'
        heading1_style.font.size = Pt(16)
        
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = '黑体'
        heading2_style.font.size = Pt(14)
    
    def _add_cover_page(self, doc: Document):
        """添加封面页"""
        # 添加标题
        title = doc.add_heading('投标文件', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加项目名称
        project_para = doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_run = project_para.add_run('项目名称：智慧城市建设项目')
        project_run.font.size = Pt(14)
        
        # 添加公司信息
        doc.add_paragraph()  # 空行
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run('投标单位：[公司名称]')
        company_run.font.size = Pt(12)
        
        # 添加日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run('投标日期：2024年12月')
        date_run.font.size = Pt(12)
        
        # 分页
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document):
        """添加目录页"""
        # 添加目录标题
        toc_title = doc.add_heading('目 录', level=1)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加目录项（简化版）
        toc_items = [
            "1. 公司介绍及资质证明",
            "2. 项目需求理解与分析",
            "3. 总体技术方案设计",
            "4. 系统架构与技术选型",
            "5. 项目实施计划与管理",
            "6. 运维服务与技术支持",
            "7. 项目预算与报价分析"
        ]
        
        for item in toc_items:
            toc_para = doc.add_paragraph(item)
            toc_para.style = 'List Number'
        
        # 分页
        doc.add_page_break()
    
    def _collect_all_sections(self) -> List[Dict[str, Any]]:
        """收集所有章节内容"""
        sections_content = []
        
        if not self.current_project_dir:
            return sections_content
        
        sections_dir = self.current_project_dir / "sections"
        if not sections_dir.exists():
            return sections_content
        
        for section_dir in sorted(sections_dir.iterdir()):
            if section_dir.is_dir():
                section_info = {
                    "title": section_dir.name,
                    "subsections": []
                }
                
                for file_item in sorted(section_dir.iterdir()):
                    if file_item.is_file() and file_item.suffix == '.md':
                        try:
                            content = file_item.read_text(encoding='utf-8')
                            has_content = "<!-- 内容待生成 -->" not in content
                            
                            section_info["subsections"].append({
                                "title": file_item.stem,
                                "content": content,
                                "has_content": has_content
                            })
                        except Exception as e:
                            logger.error(f"Failed to read {file_item}: {e}")
                
                sections_content.append(section_info)
        
        return sections_content 